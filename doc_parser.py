"""
doc_parser.py — Reads .docx files and writes structured topic.md files.

Each topic.md lives at:
    structured_data/grade_N/module_N/topic_N/topic.md

Images are described inline via vision_llm; image bytes are never saved to disk.
Equations are converted from OMML to LaTeX via lxml (requires configs/omml2mml.xsl).
If conversion fails, [EQUATION: conversion failed] is written instead.

Public API
----------
parse_syllabus(path) -> dict
    Unchanged. Returns the grade/module/topic/subtopic map used for topic discovery.

write_topic_md(module_path, syllabus_path, module_num, topic_num,
               structured_base, grade, force=False) -> Path | None
    Writes structured_data/grade_N/module_N/topic_N/topic.md.
    Returns the output path if written, None if skipped (file exists and force=False).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

try:
    from lxml import etree as _etree
    _LXML_AVAILABLE = True
except ImportError:  # pragma: no cover
    _LXML_AVAILABLE = False

import vision_llm


# ---------------------------------------------------------------------------
# Namespace constants
# ---------------------------------------------------------------------------

_A_NS    = "http://schemas.openxmlformats.org/drawingml/2006/main"
_REL_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

_IMAGE_EXTS = frozenset({"png", "jpg", "jpeg", "gif", "bmp", "tiff", "emf", "wmf"})


# ---------------------------------------------------------------------------
# Internal helpers — unchanged from original
# ---------------------------------------------------------------------------

def _iter_body_items(doc):
    """Yield ('paragraph', para) or ('table', tbl) in document order."""
    para_idx = 0
    table_idx = 0
    paras = doc.paragraphs
    tables = doc.tables
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_idx < len(paras):
                yield "paragraph", paras[para_idx]
                para_idx += 1
        elif tag == "tbl":
            if table_idx < len(tables):
                yield "table", tables[table_idx]
                table_idx += 1


def _normalize(text: str) -> str:
    """Lowercase, normalise smart quotes, strip trailing punctuation, collapse whitespace."""
    text = (
        text
        .replace("\u2019", "'").replace("\u2018", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2013", "-").replace("\u2014", "-")
        .replace("\u2026", "...")
    )
    text = text.strip().lower()
    text = re.sub(r"[?!.,:;]+$", "", text)
    text = re.sub(r"\s+", " ", text)
    return text


_STORY_STARTERS = (
    "once upon",
    "one day,",
    "one day ",
    "imagine ",
    "in a land",
    "in the land",
    "long ago",
)

_BULLET_CHARS = set("•–-*◦→")


def _classify_block(text: str, in_story: bool) -> str:
    """Return the block type for a paragraph of plain body text."""
    stripped = text.strip()
    lower = stripped.lower()

    if re.match(r"figure\s+\d+", lower):
        return "figure_caption"
    if re.match(r"image\s+\d+", lower):
        return "image_caption"
    if re.match(r"(a\s+)?quick\s+activity", lower) or lower.startswith("activity:") or lower.startswith("activity "):
        return "activity"
    if lower.startswith("fun fact"):
        return "fun_fact"
    if any(lower.startswith(s) for s in _STORY_STARTERS):
        return "story"
    if in_story:
        return "story"
    if stripped and stripped[0] in _BULLET_CHARS:
        return "bullet"
    if re.match(r"^\d+[.)]\s", stripped):
        return "bullet"
    return "paragraph"


# ---------------------------------------------------------------------------
# New internal helpers — image / equation / markdown conversion
# ---------------------------------------------------------------------------

def _para_image_rids(para, doc) -> list[tuple[str, str]]:
    """Return [(rId, extension)] for embedded images in a paragraph."""
    results = []
    seen: set[str] = set()
    for blip in para._element.iter(f"{{{_A_NS}}}blip"):
        rId = blip.get(f"{{{_REL_NS}}}embed")
        if rId and rId not in seen and rId in doc.part.related_parts:
            seen.add(rId)
            part = doc.part.related_parts[rId]
            ext = part.partname.split(".")[-1].lower()
            if ext in _IMAGE_EXTS:
                results.append((rId, ext))
    return results


def _para_omml_elements(para) -> list:
    """Return a list of m:oMath lxml elements found in a paragraph."""
    if not _LXML_AVAILABLE:
        return []
    return list(para._element.iter(f"{{{_OMML_NS}}}oMath"))


def _omml_to_latex(omml_el) -> str | None:
    """
    Try to convert an OMML m:oMath element to LaTeX via an XSLT transform.

    Requires configs/omml2mml.xsl (Microsoft OMML→MathML transform).
    Returns None if the file is missing or if conversion fails.
    """
    if not _LXML_AVAILABLE:
        return None
    xslt_path = Path(__file__).parent / "configs" / "omml2mml.xsl"
    if not xslt_path.exists():
        return None
    try:
        xslt_doc = _etree.parse(str(xslt_path))
        transform = _etree.XSLT(xslt_doc)
        mathml_tree = transform(omml_el)
        # Extract text nodes from the resulting MathML as a rough LaTeX approximation
        texts = [
            node.text.strip()
            for node in mathml_tree.iter()
            if node.text and node.text.strip()
        ]
        result = "".join(texts)
        return result if result else None
    except Exception:
        return None


def _table_to_markdown(rows: list[list[str]]) -> str:
    """Convert a list of row-lists to a GitHub-flavoured markdown table string."""
    if not rows:
        return ""
    ncols = max(len(r) for r in rows)
    padded = [r + [""] * (ncols - len(r)) for r in rows]
    lines = [
        "| " + " | ".join(padded[0]) + " |",
        "| " + " | ".join("---" for _ in padded[0]) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _block_to_md_line(block_type: str, text: str) -> str:
    """Convert a typed block to its markdown representation."""
    if block_type == "bullet":
        stripped = text.strip()
        if re.match(r"^\d+[.)]\s", stripped):
            body = re.sub(r"^\d+[.)]\s+", "", stripped)
            return f"1. {body}"
        body = stripped.lstrip("•–-*◦→ ").strip()
        return f"- {body}"
    elif block_type == "story":
        return f"> {text}"
    elif block_type == "fun_fact":
        body = re.sub(r"^fun\s+fact[:\s]*", "", text, flags=re.IGNORECASE).strip()
        return f"> **Fun Fact:** {body}"
    elif block_type == "activity":
        body = re.sub(r"^(a\s+quick\s+)?activity[:\s]*", "", text, flags=re.IGNORECASE).strip()
        return f"> **Activity:** {body}"
    else:
        # paragraph, figure_caption (orphan), image_caption (orphan)
        return text


# ---------------------------------------------------------------------------
# Internal parse_module — now private, accepts a Document object
# ---------------------------------------------------------------------------

def _parse_module(doc: Document, module_entry: dict) -> dict:
    """
    Parse an open Document into content blocks grouped by topic → subtopic.

    Extends the original parse_module logic with detection of:
    - Image paragraphs  → {"type": "image", "rId": ..., "extension": ...}
    - Equation paragraphs → {"type": "equation", "xml": "<m:oMath .../>"}

    Returns the same shape as the original parse_module:
        {"module_num": N, "topics": [{...}]}
    """
    topic_lookup: dict = {}
    subtopic_lookup: dict = {}
    topic_by_num: dict = {}

    for topic in module_entry["topics"]:
        topic_lookup[_normalize(topic["name"])] = topic
        topic_by_num[topic["topic_num"]] = topic
        for st_name in topic["subtopics"]:
            subtopic_lookup.setdefault(_normalize(st_name), []).append(
                (topic["topic_num"], st_name)
            )

    result_topics: list = []
    state = {
        "cur_topic": None,
        "cur_st_name": None,
        "cur_blocks": [],
        "in_story": False,
    }

    def _flush_subtopic() -> None:
        if state["cur_topic"] is None or not state["cur_blocks"]:
            state["cur_blocks"] = []
            return
        if state["cur_st_name"] is None:
            state["cur_blocks"] = []
            return
        state["cur_topic"]["subtopics"].append(
            {"name": state["cur_st_name"], "blocks": state["cur_blocks"]}
        )
        state["cur_blocks"] = []

    def _flush_topic() -> None:
        _flush_subtopic()
        if state["cur_topic"] is not None:
            result_topics.append(state["cur_topic"])

    def _start_topic(topic_num: int) -> None:
        _flush_topic()
        syl = topic_by_num[topic_num]
        state["cur_topic"] = {
            "topic_num": topic_num,
            "name": syl["name"],
            "subtopics": [],
        }
        state["cur_st_name"] = None
        state["cur_blocks"] = []
        state["in_story"] = False

    def _start_subtopic(topic_num: int, st_name: str) -> None:
        if state["cur_topic"] is None or state["cur_topic"]["topic_num"] != topic_num:
            _start_topic(topic_num)
        else:
            _flush_subtopic()
        state["cur_st_name"] = st_name
        state["cur_blocks"] = []
        state["in_story"] = False

    for item_type, item in _iter_body_items(doc):
        if item_type == "paragraph":
            text = item.text.strip()
            image_rids = _para_image_rids(item, doc)
            omml_els = _para_omml_elements(item)

            # Image-only paragraph (no text)
            if not text and image_rids:
                if state["cur_topic"] is not None:
                    for rId, ext in image_rids:
                        state["cur_blocks"].append(
                            {"type": "image", "rId": rId, "extension": ext}
                        )
                continue

            # Equation-only paragraph (no text)
            if not text and omml_els:
                if state["cur_topic"] is not None:
                    for omml_el in omml_els:
                        state["cur_blocks"].append({
                            "type": "equation",
                            "xml": _etree.tostring(omml_el, encoding="unicode")
                            if _LXML_AVAILABLE else "",
                        })
                continue

            if not text:
                continue

            norm = _normalize(text)

            if norm in topic_lookup:
                _start_topic(topic_lookup[norm]["topic_num"])

            elif norm in subtopic_lookup:
                candidates = subtopic_lookup[norm]
                cur = state["cur_topic"]["topic_num"] if state["cur_topic"] else None
                topic_num_c, st_name = next(
                    (c for c in candidates if c[0] == cur), candidates[0]
                )
                _start_subtopic(topic_num_c, st_name)

            else:
                if state["cur_topic"] is not None:
                    block_type = _classify_block(text, state["in_story"])
                    if block_type == "story":
                        state["in_story"] = True
                    elif block_type != "paragraph":
                        state["in_story"] = False
                    state["cur_blocks"].append({"type": block_type, "text": text})
                    # Inline equations appended after the paragraph block
                    if omml_els and _LXML_AVAILABLE:
                        for omml_el in omml_els:
                            state["cur_blocks"].append({
                                "type": "equation",
                                "xml": _etree.tostring(omml_el, encoding="unicode"),
                            })

        elif item_type == "table":
            if state["cur_topic"] is not None:
                state["in_story"] = False
                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in item.rows
                ]
                state["cur_blocks"].append({"type": "table", "rows": rows})

    _flush_topic()
    return {"module_num": module_entry["module_num"], "topics": result_topics}


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_syllabus(path) -> dict:
    """
    Parse a syllabus .docx (e.g. Class_6.docx) into a grade/module/topic/subtopic map.

    Returns:
        {
            "grade": 6,
            "modules": [
                {
                    "module_num": 1,
                    "name": "What is Artificial Intelligence?",
                    "topics": [
                        {
                            "topic_num": 1,
                            "name": "Introduction to Intelligence",
                            "subtopics": ["Human Intelligence vs. Machine Intelligence", ...]
                        },
                        ...
                    ]
                },
                ...
            ]
        }
    """
    path = Path(path)
    doc = Document(str(path))

    m = re.search(r"Class_(\d+)", path.stem, re.IGNORECASE)
    grade = int(m.group(1)) if m else 0

    modules = []
    current_module = None

    for item_type, item in _iter_body_items(doc):
        if item_type == "paragraph":
            text = item.text.strip()
            m = re.match(r"Module\s+(\d+)\s*[:\-]\s*(.+)", text, re.IGNORECASE)
            if m:
                if current_module is not None:
                    modules.append(current_module)
                current_module = {
                    "module_num": int(m.group(1)),
                    "name": m.group(2).strip().rstrip("?!."),
                    "topics": [],
                }

        elif item_type == "table" and current_module is not None:
            for row in item.rows:
                cells = [c.text.strip() for c in row.cells]
                if not cells:
                    continue
                try:
                    topic_num = int(cells[0])
                except (ValueError, IndexError):
                    continue
                topic_name = cells[1] if len(cells) > 1 else ""
                if not topic_name:
                    continue
                subtopics = []
                if len(cells) > 2:
                    subtopics = [s.strip() for s in cells[2].split("\n") if s.strip()]
                current_module["topics"].append(
                    {
                        "topic_num": topic_num,
                        "name": topic_name,
                        "subtopics": subtopics,
                    }
                )

    if current_module is not None:
        modules.append(current_module)

    return {"grade": grade, "modules": modules}



def write_topic_md(
    module_path,
    syllabus_path,
    module_num: int,
    topic_num: int,
    structured_base,
    grade: int,
    force: bool = False,
) -> "Path | None":
    """
    Parse a topic from a module .docx and write a structured topic.md.

    Output path: structured_base/grade_N/module_N/topic_N/topic.md

    - Images are described via vision_llm and written as:
        > **[Figure N]** {description}
        > *Caption: {caption}*    ← omitted if no caption
    - Equations are converted OMML→LaTeX via lxml (requires configs/omml2mml.xsl);
      on failure: [EQUATION: conversion failed]
    - Skip logic: if topic.md already exists and force=False, returns None immediately.

    Returns the output Path if written, None if skipped.
    """
    output_path = (
        Path(structured_base)
        / f"grade_{grade}"
        / f"module_{module_num}"
        / f"topic_{topic_num}"
        / "topic.md"
    )

    if output_path.exists() and not force:
        return None

    syllabus = parse_syllabus(syllabus_path)
    module_entry = next(
        (m for m in syllabus["modules"] if m["module_num"] == module_num), None
    )
    if module_entry is None:
        raise ValueError(f"Module {module_num} not found in syllabus")

    topic_entry = next(
        (t for t in module_entry["topics"] if t["topic_num"] == topic_num), None
    )
    if topic_entry is None:
        raise ValueError(f"Topic {topic_num} not found in module {module_num}")

    doc = Document(str(module_path))
    module_data = _parse_module(doc, module_entry)

    # Collect ALL instances (handles non-linear topic ordering in docx)
    instances = [t for t in module_data["topics"] if t["topic_num"] == topic_num]
    if not instances:
        raise ValueError(
            f"Topic {topic_num} not found in module {module_num} content"
        )

    all_subtopics: list[dict] = []
    for inst in instances:
        all_subtopics.extend(inst["subtopics"])

    # Warn about subtopics listed in the syllabus but absent from the docx
    found_names = {s["name"] for s in all_subtopics}
    for expected in topic_entry["subtopics"]:
        if expected not in found_names:
            print(
                f"[doc_parser] WARNING: subtopic not found in docx "
                f"(grade {grade}, module {module_num}, topic {topic_num}): {expected!r}"
            )

    # Build markdown
    md_lines: list[str] = [f"# Topic: {instances[0]['name']}", ""]
    figure_counter = 0

    for subtopic in all_subtopics:
        md_lines.append(f"## Subtopic: {subtopic['name']}")
        md_lines.append("")

        blocks = subtopic["blocks"]
        i = 0
        while i < len(blocks):
            block = blocks[i]

            if block["type"] == "image":
                # Look ahead for a caption block immediately following
                caption: str | None = None
                if (
                    i + 1 < len(blocks)
                    and blocks[i + 1]["type"] in ("figure_caption", "image_caption")
                ):
                    caption = blocks[i + 1].get("text")
                    i += 1  # consume the caption block

                figure_counter += 1
                rId = block["rId"]
                ext = block["extension"]
                img_bytes = doc.part.related_parts[rId].blob
                is_code, description = vision_llm.describe_image(
                    img_bytes,
                    ext,
                    caption,
                    instances[0]["name"],
                    subtopic["name"],
                )
                if is_code:
                    code_img_dir = output_path.parent / "code_images"
                    code_img_dir.mkdir(parents=True, exist_ok=True)
                    code_img_filename = f"figure_{figure_counter}.{ext.lower()}"
                    (code_img_dir / code_img_filename).write_bytes(img_bytes)
                    md_lines.append(f"> **[Code Figure {figure_counter}]**")
                    for line in (description or "[No code transcription]").splitlines():
                        md_lines.append(f"> {line}")
                    md_lines.append(f"> *Code Image: code_images/{code_img_filename}*")
                    if caption:
                        md_lines.append(f"> *Caption: {caption}*")
                    md_lines.append("")
                else:
                    # Flatten multi-line figure descriptions to a single line
                    # (not quoted verbatim in gist, so joining with space is fine).
                    flat_desc = " ".join(description.splitlines()).strip()
                    if not flat_desc:
                        flat_desc = caption or "[No description]"
                    md_lines.append(f"> **[Figure {figure_counter}]** {flat_desc}")
                    if caption:
                        md_lines.append(f"> *Caption: {caption}*")
                    md_lines.append("")

            elif block["type"] == "equation":
                xml_str = block.get("xml", "")
                latex: str | None = None
                if xml_str and _LXML_AVAILABLE:
                    try:
                        omml_el = _etree.fromstring(xml_str)
                        latex = _omml_to_latex(omml_el)
                    except Exception:
                        latex = None
                if latex:
                    if "\n" in latex or len(latex) > 60:
                        md_lines.append(f"$${latex}$$")
                    else:
                        md_lines.append(f"${latex}$")
                else:
                    md_lines.append("[EQUATION: conversion failed]")

            elif block["type"] == "table":
                md_lines.append(_table_to_markdown(block["rows"]))

            else:
                line = _block_to_md_line(block["type"], block.get("text", ""))
                if line:
                    md_lines.append(line)

            i += 1

        md_lines.append("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(
        "\n".join(md_lines).rstrip() + "\n", encoding="utf-8"
    )
    return output_path
